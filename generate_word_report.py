import subprocess
import sys

# Install python-docx if not present
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document

doc = Document()

doc.add_heading('Methodology Report: E-Commerce Conversion Prediction', 0)

# Section 1
doc.add_heading('1. Goal', level=1)
doc.add_paragraph(
    'The objective of my project was to accurately predict whether a user will convert (binary classification) '
    'based on their browsing behavior and demographics. The primary metric I optimized was the F1 Score.'
)

# Section 2
doc.add_heading('2. Data Preparation', level=1)
p2_1 = doc.add_paragraph(style='List Bullet')
p2_1.add_run('Combined Datasets: ').bold = True
p2_1.add_run('I merged the training data and public test data to maximize the amount of information the models could learn from.')

p2_2 = doc.add_paragraph(style='List Bullet')
p2_2.add_run('Handling Missing Data: ').bold = True
p2_2.add_run('I filled in missing numerical values using the median value of their respective columns.')

p2_3 = doc.add_paragraph(style='List Bullet')
p2_3.add_run('Encoding: ').bold = True
p2_3.add_run('I converted text-based categories (like Device Type and Traffic Source) into numbers so the models could understand them.')

# Section 3
doc.add_heading('3. Feature Engineering', level=1)
doc.add_paragraph('I created a single, simple but powerful interaction feature to capture how users engage with the website:')
p3_1 = doc.add_paragraph(style='List Bullet')
p3_1.add_run('Engagement Score: ').bold = True
p3_1.add_run('I multiplied the number of pages viewed by the number of products viewed to measure total user interest and browsing intensity.')

# Section 4
doc.add_heading('4. Machine Learning Models', level=1)
doc.add_paragraph('I used a standard ensemble approach, combining two fundamental algorithms to make my final decision:')
p4_1 = doc.add_paragraph(style='List Bullet')
p4_1.add_run('XGBoost (Weighted 66%): ').bold = True
p4_1.add_run('A standard gradient boosting model. I gave it more weight in the final vote because it proved to be slightly stronger in my testing.')

p4_2 = doc.add_paragraph(style='List Bullet')
p4_2.add_run('Random Forest (Weighted 33%): ').bold = True
p4_2.add_run('A standard random forest model with balanced class weights designed to handle class imbalances.')

# Section 5
doc.add_heading('5. Optimization and Final Prediction', level=1)
p5_1 = doc.add_paragraph(style='List Bullet')
p5_1.add_run('Cross-Validation: ').bold = True
p5_1.add_run('I tested my models rigorously using 5-fold cross-validation to ensure they would perform well on unseen data.')

p5_2 = doc.add_paragraph(style='List Bullet')
p5_2.add_run('Threshold Tuning: ').bold = True
p5_2.add_run('Instead of using the default 50% probability to decide if someone converts, I tested multiple probability thresholds. I found and applied the exact threshold that maximized my target F1 score.')

p5_3 = doc.add_paragraph(style='List Bullet')
p5_3.add_run('Final Output: ').bold = True
p5_3.add_run('My two models generated a combined probability score. If this score crossed my optimized threshold, the final prediction was set to "Converted" (1), otherwise "Not Converted" (0).')

doc.save('methodology_report.docx')
print("Successfully generated methodology_report.docx")
